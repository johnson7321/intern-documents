import os
import re
from docx import Document

# 提取中文用的正則（含簡體、繁體）
def extract_chinese(text):
    return ''.join(re.findall(r'[\u4e00-\u9fff]+', text))

# 要合併的最終文字
all_text = ""

# 處理目前資料夾下所有 .docx 檔案
for filename in os.listdir():
    if filename.endswith(".docx") and not filename.startswith("~$"):
        print(f"📄 正在處理：{filename}")
        try:
            doc = Document(filename)
        except Exception as e:
            print(f"⚠️ 無法開啟 {filename}：{e}")
            continue

        chinese_text = ""

        # 擷取段落中的文字
        for para in doc.paragraphs:
            chinese_text += extract_chinese(para.text)

        # 擷取表格中的文字（如果有）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    chinese_text += extract_chinese(cell.text)

        # 加入分隔文字與內容
        all_text += f"\n\n----- {filename} -----\n\n{chinese_text}"

# 儲存合併後的結果
output_file = "combined_chinese.txt"
with open(output_file, "w", encoding="utf-8") as f:
    f.write(all_text)

print(f"\n✅ 合併完成！中文文字已儲存至：{output_file}")
